from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# === TITLE PAGE ===
for _ in range(4):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("RAJALAKSHMI ENGINEERING COLLEGE")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(75, 0, 130)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("AI-Generated Learning Video Script")
r2.bold = True; r2.font.size = Pt(18); r2.font.color.rgb = RGBColor(100, 50, 150)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("Operating Systems — Explained by Aadhi 🦌")
r3.font.size = Pt(16); r3.font.color.rgb = RGBColor(50, 50, 50)

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run("Total Duration: 5 Minutes (30 Scenes × 10 Seconds Each)")
r4.font.size = Pt(13); r4.font.color.rgb = RGBColor(80, 80, 80)

t5 = doc.add_paragraph()
t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = t5.add_run("Format: Minimalistic Lecture-Style Slideshow")
r5.font.size = Pt(12); r5.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# === MASCOT DESCRIPTION PAGE ===
h = doc.add_heading("Mascot Reference: Aadhi", level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(75, 0, 130)

doc.add_paragraph(
    "Aadhi is an energetic, anthropomorphic antelope (blackbuck) exuding school spirit and confidence. "
    "Key visual features:"
)
bullets = [
    "Striking dark brown spiraled horns",
    "Two-toned tan and dark brown fur",
    "Sleek purple wraparound sunglasses — sporty, modern edge",
    "Vibrant purple zip-up track jacket with yellow trim",
    "Large yellow 'R' enclosing a lit torch emblem on chest (Rajalakshmi logo)",
    "Self-assured smirk expression",
    "Default pose: right fist raised triumphantly, left hand on hip",
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()
h2 = doc.add_heading("Global Visual & Continuity Rules", level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(75, 0, 130)

rules = [
    "Background: Clean, minimalistic gradient (light purple-to-white or soft grey). No realistic scenarios.",
    "Layout: Lecture-style slideshow — clear heading at top, bullet points / diagrams in center, Aadhi on the right side.",
    "Aadhi's position: Consistently on the RIGHT side of the frame, roughly 30% width, gesturing toward the content.",
    "Text style: Bold headings (purple), clean sans-serif body text (dark grey/black), key terms highlighted in yellow.",
    "Transitions: Smooth fade/slide between scenes to ensure seamless merging.",
    "Aadhi's appearance must remain IDENTICAL across all 30 scenes (same outfit, horns, sunglasses, proportions).",
    "Aspect ratio: 16:9 widescreen.",
    "Aadhi speaks directly to the viewer (slight head tilt, mouth animated to narration).",
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# === SCENES ===
scenes = [
    # --- INTRODUCTION (Scenes 1-2) ---
    {
        "num": 1, "section": "INTRODUCTION", "time": "0:00 – 0:10",
        "heading_on_screen": "Welcome to Operating Systems!",
        "visual": (
            "Clean purple gradient background. The Rajalakshmi Engineering College logo fades in at the top center. "
            "Aadhi slides in from the right with a wave animation, standing confidently with his right fist raised. "
            "Title text animates in: \"OPERATING SYSTEMS\" in large bold purple letters, with subtitle \"Explained by Aadhi\" below."
        ),
        "narration": (
            "Hey there, REC students! I'm Aadhi, your campus mascot — and today, I'm going to break down "
            "Operating Systems for you in just five minutes! Let's get started!"
        ),
        "on_screen_text": "OPERATING SYSTEMS\nExplained by Aadhi | Rajalakshmi Engineering College",
        "notes": "Opening scene. Aadhi is energetic, welcoming. Sets the tone for the entire video."
    },
    {
        "num": 2, "section": "INTRODUCTION", "time": "0:10 – 0:20",
        "heading_on_screen": "What We'll Cover Today",
        "visual": (
            "Aadhi gestures toward a vertical list that appears one-by-one on the left side of the screen. "
            "Each topic appears with a small icon next to it. Aadhi points to each item as it appears. "
            "Background remains the clean gradient."
        ),
        "narration": (
            "Here's our roadmap! We'll cover what an OS is, its functions, process management, memory, "
            "file systems, CPU scheduling, deadlocks, synchronization, and types of OS. Let's dive in!"
        ),
        "on_screen_text": (
            "Today's Agenda:\n• What is an OS?\n• Functions of OS\n• Process Management\n"
            "• Memory Management\n• File Systems\n• CPU Scheduling\n• Deadlocks\n• Synchronization\n• Types of OS"
        ),
        "notes": "Roadmap scene. Creates anticipation. Smooth transition to Scene 3."
    },

    # --- WHAT IS AN OPERATING SYSTEM (Scenes 3-5) ---
    {
        "num": 3, "section": "WHAT IS AN OPERATING SYSTEM", "time": "0:20 – 0:30",
        "heading_on_screen": "What is an Operating System?",
        "visual": (
            "Heading slides in at top: \"What is an Operating System?\" in bold purple. "
            "Center shows a simple layered diagram: USER on top → APPLICATION in middle → OPERATING SYSTEM → HARDWARE at bottom. "
            "Aadhi on the right, pointing at the OS layer with an explanatory gesture."
        ),
        "narration": (
            "So, what exactly is an Operating System? It's the software that sits between YOU and the hardware. "
            "It manages everything — from your apps to the CPU and memory!"
        ),
        "on_screen_text": "An OS is system software that manages hardware and software resources.",
        "notes": "Introduces core definition. Layered diagram is key visual."
    },
    {
        "num": 4, "section": "WHAT IS AN OPERATING SYSTEM", "time": "0:30 – 0:40",
        "heading_on_screen": "OS as an Intermediary",
        "visual": (
            "Same heading persists. The layered diagram transforms into an animated flow: "
            "arrows flow from User → App → OS → Hardware and back. "
            "Aadhi traces the flow with his finger, nodding approvingly."
        ),
        "narration": (
            "Think of the OS as a translator! When you click an app, the OS talks to the hardware to make it happen. "
            "Without it, your computer is just a pile of circuits!"
        ),
        "on_screen_text": "OS = Intermediary between User and Hardware\n→ Translates requests into hardware actions",
        "notes": "Deepens understanding with analogy. Arrow animation shows data flow."
    },
    {
        "num": 5, "section": "WHAT IS AN OPERATING SYSTEM", "time": "0:40 – 0:50",
        "heading_on_screen": "Examples of Operating Systems",
        "visual": (
            "Heading: \"Examples of Operating Systems\". Below, four clean icons appear in a row: "
            "Windows logo, Apple/macOS logo, Linux penguin, Android robot. "
            "Each label appears below its icon. Aadhi gives a thumbs-up."
        ),
        "narration": (
            "You use an OS every day! Windows, macOS, Linux, Android — they're all operating systems. "
            "Each one manages your device differently, but the core idea is the same!"
        ),
        "on_screen_text": "Windows | macOS | Linux | Android",
        "notes": "Relatable examples. Transition to Functions section next."
    },

    # --- FUNCTIONS OF OS (Scenes 6-8) ---
    {
        "num": 6, "section": "FUNCTIONS OF OS", "time": "0:50 – 1:00",
        "heading_on_screen": "Key Functions of an OS",
        "visual": (
            "New section slide. Heading: \"Key Functions of an OS\" in bold purple. "
            "A central hub diagram appears with \"OS\" in the middle circle, and 5 branches radiating outward: "
            "Process Mgmt, Memory Mgmt, File Mgmt, I/O Mgmt, Security. Aadhi stands beside it, arms open in a presenting gesture."
        ),
        "narration": (
            "An OS does FIVE major things — it manages processes, memory, files, input-output devices, and security. "
            "Let's see each one!"
        ),
        "on_screen_text": "5 Key Functions:\n1. Process Management\n2. Memory Management\n3. File Management\n4. I/O Management\n5. Security & Protection",
        "notes": "Overview of all functions. Hub diagram is the key visual."
    },
    {
        "num": 7, "section": "FUNCTIONS OF OS", "time": "1:00 – 1:10",
        "heading_on_screen": "Resource Allocation & I/O Management",
        "visual": (
            "Heading stays. Left side shows a simple animation: CPU, RAM, and Disk icons with arrows showing the OS distributing resources "
            "to multiple app icons. Aadhi points at the distribution arrows."
        ),
        "narration": (
            "Resource allocation means the OS decides which app gets how much CPU, memory, and disk. "
            "It also handles all your input-output — keyboard, mouse, display, everything!"
        ),
        "on_screen_text": "Resource Allocation: CPU, RAM, Disk → distributed to apps\nI/O Management: Keyboard, Mouse, Display",
        "notes": "Shows resource distribution visually."
    },
    {
        "num": 8, "section": "FUNCTIONS OF OS", "time": "1:10 – 1:20",
        "heading_on_screen": "Security & Protection",
        "visual": (
            "Heading: \"Security & Protection\". A shield icon appears center-left with a lock symbol. "
            "Around it: \"Authentication\", \"Access Control\", \"Encryption\" labels fade in. "
            "Aadhi crosses his arms confidently, nodding."
        ),
        "narration": (
            "Security is crucial! The OS uses authentication — like passwords — access control to restrict files, "
            "and encryption to protect your data. Your OS is your digital bodyguard!"
        ),
        "on_screen_text": "Security Functions:\n• Authentication (passwords, biometrics)\n• Access Control (permissions)\n• Encryption (data protection)",
        "notes": "Completes Functions section. Transitions to Process Management."
    },

    # --- PROCESS MANAGEMENT (Scenes 9-12) ---
    {
        "num": 9, "section": "PROCESS MANAGEMENT", "time": "1:20 – 1:30",
        "heading_on_screen": "What is a Process?",
        "visual": (
            "New section. Heading: \"Process Management\" with sub-heading \"What is a Process?\". "
            "Center-left: a simple comparison — a document icon labeled \"Program (passive)\" with an arrow pointing to a running gear icon "
            "labeled \"Process (active)\". Aadhi points at the arrow."
        ),
        "narration": (
            "A program sitting on your disk is just code — it's passive. But the moment you run it, it becomes a PROCESS — "
            "an active entity with its own memory and state!"
        ),
        "on_screen_text": "Program = Passive code on disk\nProcess = Program in execution (active)",
        "notes": "Key distinction between program and process."
    },
    {
        "num": 10, "section": "PROCESS MANAGEMENT", "time": "1:30 – 1:40",
        "heading_on_screen": "Process States",
        "visual": (
            "Heading: \"Process States\". A state diagram appears: five colored boxes — New → Ready → Running → Waiting → Terminated, "
            "connected by directional arrows. Each state is a different color. Aadhi traces the flow with his hand."
        ),
        "narration": (
            "Every process goes through states: New, Ready, Running, Waiting, and Terminated. "
            "The OS moves processes between these states to keep everything running smoothly!"
        ),
        "on_screen_text": "Process States:\nNew → Ready → Running → Waiting → Terminated",
        "notes": "State transition diagram is the key visual."
    },
    {
        "num": 11, "section": "PROCESS MANAGEMENT", "time": "1:40 – 1:50",
        "heading_on_screen": "Process Control Block (PCB)",
        "visual": (
            "Heading: \"Process Control Block\". A rectangular card/table appears showing fields: "
            "Process ID, State, Program Counter, CPU Registers, Memory Info. "
            "Aadhi holds up his hand as if presenting a card."
        ),
        "narration": (
            "Each process has an ID card called the PCB — Process Control Block. It stores the process ID, "
            "current state, program counter, registers, and memory info. The OS uses this to track every process!"
        ),
        "on_screen_text": "PCB Contains:\n• Process ID\n• Process State\n• Program Counter\n• CPU Registers\n• Memory Information",
        "notes": "PCB explained with a visual card metaphor."
    },
    {
        "num": 12, "section": "PROCESS MANAGEMENT", "time": "1:50 – 2:00",
        "heading_on_screen": "Context Switching",
        "visual": (
            "Heading: \"Context Switching\". Animation shows two process boxes (P1 and P2) with a CPU in the center. "
            "Arrows show P1 saving its state, CPU switching to P2, P2 loading its state. "
            "Aadhi mimics a switching motion with his hands."
        ),
        "narration": (
            "Context switching is when the CPU saves one process's state and loads another. "
            "It happens so fast, you think everything runs at once — but the CPU is actually juggling!"
        ),
        "on_screen_text": "Context Switch:\nSave state of P1 → Load state of P2\n(Enables multitasking!)",
        "notes": "Completes Process Management. Smooth transition to Memory Management."
    },

    # --- MEMORY MANAGEMENT (Scenes 13-16) ---
    {
        "num": 13, "section": "MEMORY MANAGEMENT", "time": "2:00 – 2:10",
        "heading_on_screen": "Memory Management",
        "visual": (
            "New section. Heading: \"Memory Management\". Center shows a simplified RAM stick icon with labeled sections: "
            "OS area, Process 1, Process 2, Free space. Aadhi gestures at the RAM diagram."
        ),
        "narration": (
            "Memory management is how the OS allocates and tracks RAM. It decides which process gets how much memory "
            "and makes sure they don't step on each other's toes!"
        ),
        "on_screen_text": "Memory Management:\n• Allocates RAM to processes\n• Tracks usage\n• Prevents conflicts",
        "notes": "Introduces memory management with RAM visualization."
    },
    {
        "num": 14, "section": "MEMORY MANAGEMENT", "time": "2:10 – 2:20",
        "heading_on_screen": "Contiguous vs Non-Contiguous Allocation",
        "visual": (
            "Heading: \"Memory Allocation\". Two side-by-side diagrams: Left — contiguous blocks (neat, stacked), "
            "Right — non-contiguous blocks (scattered, linked). Labels below each. Aadhi points left then right."
        ),
        "narration": (
            "Memory can be allocated in two ways — contiguous, where each process gets one continuous block, "
            "or non-contiguous, where it's split into pieces scattered across RAM!"
        ),
        "on_screen_text": "Contiguous: One continuous block per process\nNon-Contiguous: Split into pages/segments",
        "notes": "Side-by-side comparison for clarity."
    },
    {
        "num": 15, "section": "MEMORY MANAGEMENT", "time": "2:20 – 2:30",
        "heading_on_screen": "Paging",
        "visual": (
            "Heading: \"Paging\". Diagram shows: Process divided into equal \"pages\" on left, "
            "RAM divided into equal \"frames\" on right, with mapping arrows between them. "
            "A small page table is shown. Aadhi points at the mapping."
        ),
        "narration": (
            "Paging divides memory into fixed-size blocks — pages for the process and frames for RAM. "
            "A page table maps which page goes into which frame. No external fragmentation!"
        ),
        "on_screen_text": "Paging:\n• Process → Pages (fixed size)\n• RAM → Frames (fixed size)\n• Page Table maps Pages to Frames\n• No external fragmentation",
        "notes": "Core paging concept with mapping visual."
    },
    {
        "num": 16, "section": "MEMORY MANAGEMENT", "time": "2:30 – 2:40",
        "heading_on_screen": "Virtual Memory",
        "visual": (
            "Heading: \"Virtual Memory\". Diagram shows a large virtual address space on the left (bigger) "
            "mapping to a smaller physical RAM on the right, with a hard disk below labeled \"Swap Space\". "
            "Arrows show pages swapping in and out. Aadhi looks impressed, pointing at the swap."
        ),
        "narration": (
            "Virtual memory lets you run programs larger than your actual RAM! "
            "The OS swaps pages between RAM and disk, creating an illusion of unlimited memory. Clever, right?"
        ),
        "on_screen_text": "Virtual Memory:\n• Larger than physical RAM\n• Uses disk as swap space\n• Pages swapped in/out as needed",
        "notes": "Completes Memory Management. Transitions to File System."
    },

    # --- FILE SYSTEM (Scenes 17-19) ---
    {
        "num": 17, "section": "FILE SYSTEM", "time": "2:40 – 2:50",
        "heading_on_screen": "File System",
        "visual": (
            "New section. Heading: \"File System\". A tree structure appears: Root folder at top, "
            "branching into subfolders (Documents, Programs, System), each with files inside. "
            "Aadhi stands beside the tree, presenting it."
        ),
        "narration": (
            "The file system organizes your data on disk! It uses a hierarchical tree structure — "
            "folders within folders — so you can store, find, and manage your files easily."
        ),
        "on_screen_text": "File System:\n• Organizes data on disk\n• Hierarchical tree structure\n• Directories and files",
        "notes": "File system introduction with tree visual."
    },
    {
        "num": 18, "section": "FILE SYSTEM", "time": "2:50 – 3:00",
        "heading_on_screen": "File Operations & Access Methods",
        "visual": (
            "Heading: \"File Operations\". A list of operations appears with small icons: "
            "Create, Read, Write, Delete, Open, Close. Below, two access methods: Sequential and Direct. "
            "Aadhi counts them off on his fingers."
        ),
        "narration": (
            "The OS provides operations like create, read, write, and delete for files. "
            "Access can be sequential — reading one record after another — or direct, jumping to any position!"
        ),
        "on_screen_text": "Operations: Create | Read | Write | Delete | Open | Close\nAccess: Sequential vs Direct (Random)",
        "notes": "File operations and access methods."
    },
    {
        "num": 19, "section": "FILE SYSTEM", "time": "3:00 – 3:10",
        "heading_on_screen": "Disk Allocation Methods",
        "visual": (
            "Heading: \"Disk Allocation\". Three small diagrams side by side: "
            "1) Contiguous blocks, 2) Linked list blocks with pointers, 3) Indexed with an index block. "
            "Labels below each. Aadhi points to each one."
        ),
        "narration": (
            "Files are stored on disk using three methods — Contiguous allocation, Linked allocation with pointers, "
            "and Indexed allocation using an index block. Each has its trade-offs!"
        ),
        "on_screen_text": "Allocation Methods:\n1. Contiguous\n2. Linked\n3. Indexed",
        "notes": "Completes File System. Transitions to CPU Scheduling."
    },

    # --- CPU SCHEDULING (Scenes 20-22) ---
    {
        "num": 20, "section": "CPU SCHEDULING", "time": "3:10 – 3:20",
        "heading_on_screen": "CPU Scheduling",
        "visual": (
            "New section. Heading: \"CPU Scheduling\". Center shows a CPU icon with a queue of processes (P1, P2, P3, P4) "
            "waiting in line. An arrow shows P1 entering the CPU. Aadhi gestures like a traffic controller."
        ),
        "narration": (
            "CPU scheduling decides which process gets the CPU and when. "
            "Multiple processes compete for the CPU, and the scheduler picks the winner!"
        ),
        "on_screen_text": "CPU Scheduling:\n• Decides which process runs next\n• Maximizes CPU utilization\n• Managed by the Scheduler",
        "notes": "CPU scheduling introduction with queue visual."
    },
    {
        "num": 21, "section": "CPU SCHEDULING", "time": "3:20 – 3:30",
        "heading_on_screen": "Scheduling Algorithms",
        "visual": (
            "Heading: \"Scheduling Algorithms\". A table/list appears: "
            "FCFS, SJF, Priority, Round Robin — each with a one-line description. "
            "Round Robin is highlighted in yellow as the most common. Aadhi points at Round Robin."
        ),
        "narration": (
            "There are several algorithms — First Come First Serve, Shortest Job First, Priority scheduling, "
            "and Round Robin, which gives each process a time slice. Round Robin is super popular!"
        ),
        "on_screen_text": "Algorithms:\n• FCFS — First in, first served\n• SJF — Shortest job first\n• Priority — Highest priority first\n• Round Robin — Time slice for each (★)",
        "notes": "Key algorithms listed. Round Robin highlighted."
    },
    {
        "num": 22, "section": "CPU SCHEDULING", "time": "3:30 – 3:40",
        "heading_on_screen": "Preemptive vs Non-Preemptive",
        "visual": (
            "Heading: \"Preemptive vs Non-Preemptive\". Two side-by-side panels: "
            "Left — \"Preemptive\" with a process being interrupted mid-execution, "
            "Right — \"Non-Preemptive\" with a process running to completion. Aadhi makes a comparison gesture."
        ),
        "narration": (
            "In preemptive scheduling, the OS can interrupt a running process for a higher-priority one. "
            "In non-preemptive, the process runs until it finishes or voluntarily gives up the CPU!"
        ),
        "on_screen_text": "Preemptive: OS can interrupt processes\nNon-Preemptive: Process runs to completion",
        "notes": "Completes CPU Scheduling. Transitions to Deadlock."
    },

    # --- DEADLOCK (Scenes 23-25) ---
    {
        "num": 23, "section": "DEADLOCK", "time": "3:40 – 3:50",
        "heading_on_screen": "What is a Deadlock?",
        "visual": (
            "New section. Heading: \"Deadlock\". Center shows two processes (P1 and P2) each holding one resource "
            "and waiting for the other's resource — forming a circular wait illustrated with crossed arrows. "
            "Aadhi looks worried, hands on his head."
        ),
        "narration": (
            "Deadlock is a nightmare scenario! Two or more processes are stuck, each waiting for a resource "
            "the other holds. Nobody can move — it's a total traffic jam!"
        ),
        "on_screen_text": "Deadlock: Two or more processes waiting for each other's resources\n→ No progress possible!",
        "notes": "Deadlock introduction with circular wait visual."
    },
    {
        "num": 24, "section": "DEADLOCK", "time": "3:50 – 4:00",
        "heading_on_screen": "Conditions for Deadlock",
        "visual": (
            "Heading: \"4 Conditions for Deadlock\". Four boxes appear one by one: "
            "1. Mutual Exclusion, 2. Hold & Wait, 3. No Preemption, 4. Circular Wait. "
            "All four must be true — shown by \"AND\" connectors. Aadhi raises four fingers."
        ),
        "narration": (
            "Deadlock needs ALL four conditions: Mutual Exclusion — only one can use a resource, "
            "Hold and Wait, No Preemption, and Circular Wait. Break any one, and deadlock is avoided!"
        ),
        "on_screen_text": "4 Necessary Conditions (ALL must hold):\n1. Mutual Exclusion\n2. Hold & Wait\n3. No Preemption\n4. Circular Wait",
        "notes": "Four conditions clearly listed."
    },
    {
        "num": 25, "section": "DEADLOCK", "time": "4:00 – 4:10",
        "heading_on_screen": "Handling Deadlocks",
        "visual": (
            "Heading: \"Handling Deadlocks\". Three strategies appear as cards: "
            "\"Prevention\" (shield icon), \"Avoidance — Banker's Algorithm\" (bank icon), "
            "\"Detection & Recovery\" (magnifying glass icon). Aadhi gives a thumbs-up."
        ),
        "narration": (
            "We handle deadlocks in three ways — Prevention by breaking conditions, "
            "Avoidance using the Banker's Algorithm, or Detection and Recovery after it happens!"
        ),
        "on_screen_text": "Strategies:\n• Prevention — Break one of 4 conditions\n• Avoidance — Banker's Algorithm\n• Detection & Recovery",
        "notes": "Completes Deadlock. Transitions to Synchronization."
    },

    # --- SYNCHRONIZATION (Scenes 26-27) ---
    {
        "num": 26, "section": "SYNCHRONIZATION", "time": "4:10 – 4:20",
        "heading_on_screen": "Process Synchronization",
        "visual": (
            "New section. Heading: \"Process Synchronization\". Center shows two processes trying to access "
            "the same shared variable — a \"Critical Section\" box highlighted in red. "
            "Arrows show conflict. Aadhi holds up a stop sign gesture."
        ),
        "narration": (
            "When multiple processes access shared data, chaos can happen! "
            "Synchronization ensures only one process enters the critical section at a time — preventing race conditions!"
        ),
        "on_screen_text": "Synchronization:\n• Protects shared resources\n• Critical Section — only one process at a time\n• Prevents race conditions",
        "notes": "Introduces synchronization and critical section."
    },
    {
        "num": 27, "section": "SYNCHRONIZATION", "time": "4:20 – 4:30",
        "heading_on_screen": "Synchronization Tools",
        "visual": (
            "Heading: \"Synchronization Tools\". Three items appear: "
            "\"Mutex\" (lock icon), \"Semaphore\" (traffic signal icon with counter), \"Monitor\" (enclosed box icon). "
            "Brief description under each. Aadhi points at each tool."
        ),
        "narration": (
            "The key tools are — Mutex, a simple lock for one process, "
            "Semaphores with a counter for multiple resources, and Monitors that combine locking with condition variables!"
        ),
        "on_screen_text": "Tools:\n• Mutex — Binary lock (0 or 1)\n• Semaphore — Counter-based signaling\n• Monitor — High-level synchronization construct",
        "notes": "Completes Synchronization. Transitions to Types of OS."
    },

    # --- TYPES OF OS (Scenes 28-29) ---
    {
        "num": 28, "section": "TYPES OF OS", "time": "4:30 – 4:40",
        "heading_on_screen": "Types of Operating Systems",
        "visual": (
            "New section. Heading: \"Types of Operating Systems\". "
            "Four quadrants appear: Batch OS (stack of punch cards icon), Time-Sharing OS (clock with users), "
            "Real-Time OS (speedometer icon), Distributed OS (networked computers icon). "
            "Aadhi sweeps his hand across all four."
        ),
        "narration": (
            "There are several types of OS! Batch OS processes jobs in groups, "
            "Time-Sharing lets multiple users share the CPU, Real-Time OS handles time-critical tasks, "
            "and Distributed OS spans multiple machines!"
        ),
        "on_screen_text": "Types:\n• Batch OS\n• Time-Sharing OS\n• Real-Time OS (RTOS)\n• Distributed OS",
        "notes": "Overview of OS types with quadrant visual."
    },
    {
        "num": 29, "section": "TYPES OF OS", "time": "4:40 – 4:50",
        "heading_on_screen": "Modern OS Examples",
        "visual": (
            "Heading: \"Modern OS in Action\". A comparison table shows: "
            "Desktop → Windows/macOS/Linux, Mobile → Android/iOS, Embedded → RTOS, Cloud → Distributed. "
            "Each row has an icon. Aadhi leans in and points at the table."
        ),
        "narration": (
            "In the real world — desktops run Windows, macOS, or Linux; phones run Android or iOS; "
            "embedded devices use RTOS; and cloud infrastructure uses distributed systems. OS is everywhere!"
        ),
        "on_screen_text": "Desktop: Windows, macOS, Linux\nMobile: Android, iOS\nEmbedded: RTOS\nCloud: Distributed OS",
        "notes": "Connects theory to real-world. Transitions to conclusion."
    },

    # --- CONCLUSION (Scene 30) ---
    {
        "num": 30, "section": "CONCLUSION", "time": "4:50 – 5:00",
        "heading_on_screen": "That's a Wrap! 🎉",
        "visual": (
            "Final scene. Heading: \"That's a Wrap!\" with confetti animation. "
            "A quick recap list flashes on the left: all topics covered with checkmarks. "
            "Aadhi returns to his signature triumphant pose — right fist raised high, left hand on hip, "
            "big confident smirk. Rajalakshmi logo fades in at the bottom. "
            "Text: \"All the best, REC students!\" appears below Aadhi."
        ),
        "narration": (
            "And that's Operating Systems in five minutes! We covered the OS, its functions, "
            "processes, memory, file systems, scheduling, deadlocks, synchronization, and OS types. "
            "You've got this, REC! Aadhi believes in you — go ace that exam! 💪"
        ),
        "on_screen_text": "✅ OS Basics ✅ Functions ✅ Processes ✅ Memory\n✅ File Systems ✅ Scheduling ✅ Deadlocks\n✅ Synchronization ✅ Types of OS\n\nAll the best, REC students! — Aadhi 🦌",
        "notes": "Closing scene. Triumphant pose. Rajalakshmi branding. Motivational ending."
    },
]

# Write scenes
for i, scene in enumerate(scenes):
    # Section header (only for first scene of each section)
    if i == 0 or scenes[i-1]["section"] != scene["section"]:
        sh = doc.add_heading(f'Section: {scene["section"]}', level=1)
        for run in sh.runs:
            run.font.color.rgb = RGBColor(75, 0, 130)

    # Scene heading
    sh2 = doc.add_heading(f'Scene {scene["num"]} — {scene["heading_on_screen"]}', level=2)
    for run in sh2.runs:
        run.font.color.rgb = RGBColor(100, 50, 150)

    # Time
    p_time = doc.add_paragraph()
    r_time_label = p_time.add_run("⏱ Timestamp: ")
    r_time_label.bold = True
    p_time.add_run(scene["time"])

    # Visual description
    p_vis_label = doc.add_paragraph()
    r_vis_label = p_vis_label.add_run("🎬 Visual Description (Gemini Video Prompt):")
    r_vis_label.bold = True
    r_vis_label.font.color.rgb = RGBColor(0, 100, 0)

    p_vis = doc.add_paragraph(scene["visual"])
    p_vis.paragraph_format.left_indent = Cm(1)

    # Narration
    p_nar_label = doc.add_paragraph()
    r_nar_label = p_nar_label.add_run("🎙 Narration (Aadhi's Dialogue):")
    r_nar_label.bold = True
    r_nar_label.font.color.rgb = RGBColor(0, 0, 150)

    p_nar = doc.add_paragraph()
    p_nar.paragraph_format.left_indent = Cm(1)
    r_nar = p_nar.add_run(f'"{scene["narration"]}"')
    r_nar.italic = True

    # On-screen text
    p_ost_label = doc.add_paragraph()
    r_ost_label = p_ost_label.add_run("📝 On-Screen Text:")
    r_ost_label.bold = True
    r_ost_label.font.color.rgb = RGBColor(150, 80, 0)

    p_ost = doc.add_paragraph(scene["on_screen_text"])
    p_ost.paragraph_format.left_indent = Cm(1)

    # Continuity notes
    p_notes_label = doc.add_paragraph()
    r_notes_label = p_notes_label.add_run("🔗 Continuity Notes: ")
    r_notes_label.bold = True
    p_notes_label.add_run(scene["notes"])

    # Separator
    if scene["num"] < 30:
        sep = doc.add_paragraph()
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sep = sep.add_run("— — — — — — — — — — — — — — —")
        r_sep.font.color.rgb = RGBColor(180, 180, 180)

# === APPENDIX ===
doc.add_page_break()
ah = doc.add_heading("Appendix: Gemini Video Generation Tips", level=1)
for run in ah.runs:
    run.font.color.rgb = RGBColor(75, 0, 130)

tips = [
    "Use each scene's Visual Description as the primary prompt for Gemini video generation.",
    "Prepend each prompt with the mascot description to maintain Aadhi's consistent appearance.",
    "Add 'minimalistic lecture-style background, clean purple gradient, 16:9 aspect ratio' to every prompt.",
    "For continuity: ensure Aadhi is ALWAYS on the right side, same outfit, same sunglasses, same proportions.",
    "Generate scenes sequentially and review each for consistency before moving to the next.",
    "Use the narration text to generate AI voice-over (e.g., using ElevenLabs, Google TTS, or similar).",
    "Merge all 30 clips in order using a video editor (e.g., CapCut, DaVinci Resolve, or ffmpeg).",
    "Add smooth fade transitions (0.3-0.5s) between scenes for seamless merging.",
    "Overlay the on-screen text using the video editor's text tool for crisp, readable typography.",
    "Final export: 1080p, 30fps, H.264 codec for maximum compatibility.",
]
for tip in tips:
    doc.add_paragraph(tip, style='List Bullet')

# Save
output_path = "/home/prasaz_nat/desk/projects/kutty-koncepts/Aadhi_OS_Video_Script.docx"
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
